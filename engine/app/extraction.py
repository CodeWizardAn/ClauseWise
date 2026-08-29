"""Text extraction for the ClauseWise engine sidecar.

This module does exactly one thing: turn document bytes into plain text.
No auth, no database, no business logic, no network calls — those belong to
the Next.js app (see IMPLEMENTATION.md section 3).

Design rules enforced here:
  * Bytes are processed in memory. The sidecar never writes an uploaded file
    to disk, so there is no raw file left behind to leak or to clean up.
  * Failures are loud. If a document cannot be read, or reads as empty, we
    raise with a human-readable reason instead of returning placeholder text.
  * Nothing is silently truncated. Every page and paragraph we can read is
    returned, and the page count we report is the real one.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field

import pymupdf
from docx import Document as DocxDocument

# Extension -> the kind of extractor we dispatch to.
SUPPORTED_EXTENSIONS: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
}

# Tried in order for plain-text files. utf-8-sig strips a BOM if one is present.
TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "cp1252", "latin-1")


class ExtractionError(Exception):
    """Raised when a document cannot be turned into usable text.

    `status` is the HTTP status the API layer should surface, so the reason a
    document failed reaches the user's screen instead of being swallowed.
    """

    def __init__(self, message: str, status: int = 422) -> None:
        super().__init__(message)
        self.message = message
        self.status = status


@dataclass
class ExtractionResult:
    text: str
    kind: str
    pages: int | None = None
    notes: list[str] = field(default_factory=list)


def detect_kind(filename: str) -> str:
    """Map a filename to an extractor, or fail loudly if we do not support it."""
    lowered = filename.lower()
    for extension, kind in SUPPORTED_EXTENSIONS.items():
        if lowered.endswith(extension):
            return kind
    supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
    raise ExtractionError(
        f"Unsupported file type. This build accepts {supported} only.",
        status=415,
    )


def extract(data: bytes, filename: str) -> ExtractionResult:
    """Extract text from document bytes, dispatching on the file extension."""
    kind = detect_kind(filename)
    if not data:
        raise ExtractionError("The uploaded file is empty (0 bytes).", status=400)

    if kind == "pdf":
        result = _extract_pdf(data)
    elif kind == "docx":
        result = _extract_docx(data)
    else:
        result = _extract_txt(data)

    if not result.text.strip():
        raise ExtractionError(_empty_reason(kind))
    return result


def _empty_reason(kind: str) -> str:
    if kind == "pdf":
        return (
            "No text layer found in this PDF. It is most likely a scan or a photo. "
            "OCR is not part of this build, so the document cannot be read yet."
        )
    return "The document was read successfully but contains no text."


def _extract_pdf(data: bytes) -> ExtractionResult:
    try:
        document = pymupdf.open(stream=data, filetype="pdf")
    except Exception as error:  # noqa: BLE001 - surface the real reason
        raise ExtractionError(f"This PDF could not be opened: {error}") from error

    with document:
        if document.needs_pass:
            raise ExtractionError(
                "This PDF is password protected, so its text cannot be read.",
                status=400,
            )
        page_count = document.page_count
        pages = [page.get_text("text") for page in document]

    empty_pages = sum(1 for page in pages if not page.strip())
    notes: list[str] = []
    if empty_pages:
        # Honest reporting: say which pages had no text layer rather than
        # quietly returning a shorter document than the user handed us.
        notes.append(
            f"{empty_pages} of {page_count} pages had no text layer "
            "(likely scanned images) and contributed no text."
        )

    return ExtractionResult(
        text="\n\n".join(pages).strip(),
        kind="pdf",
        pages=page_count,
        notes=notes,
    )


def _extract_docx(data: bytes) -> ExtractionResult:
    try:
        document = DocxDocument(io.BytesIO(data))
    except Exception as error:  # noqa: BLE001 - surface the real reason
        raise ExtractionError(f"This DOCX could not be opened: {error}") from error

    blocks = [paragraph.text.strip() for paragraph in document.paragraphs]

    # Rental and loan agreements routinely put the schedule of rent, deposit
    # and instalments in a table. Skipping tables would silently drop exactly
    # the figures the later phases need, so read them too.
    table_count = 0
    for table in document.tables:
        table_count += 1
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                blocks.append(line)

    notes: list[str] = []
    if table_count:
        notes.append(f"Read {table_count} table(s) in addition to the body text.")

    return ExtractionResult(
        text="\n\n".join(block for block in blocks if block).strip(),
        kind="docx",
        notes=notes,
    )


def _extract_txt(data: bytes) -> ExtractionResult:
    for encoding in TEXT_ENCODINGS:
        try:
            decoded = data.decode(encoding)
        except UnicodeDecodeError:
            continue
        notes = []
        if encoding not in ("utf-8-sig", "utf-8"):
            notes.append(f"File was not valid UTF-8; decoded as {encoding}.")
        return ExtractionResult(text=decoded.strip(), kind="txt", notes=notes)

    attempted = ", ".join(TEXT_ENCODINGS)
    raise ExtractionError(
        f"This text file could not be decoded (tried {attempted}).",
        status=400,
    )
